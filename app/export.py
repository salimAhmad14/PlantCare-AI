"""PDF and Word export.

Both formats carry the same content in the same order: diagnosis, action plan,
the evidence blocks with their citations, caveats, sources. A farmer printing the
PDF and a student pasting the DOCX into a report should be reading the same thing.

Images come out of the report as `data:image/...;base64,...` URIs, so they are
decoded back to bytes here rather than re-read from disk.
"""

import base64
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (HRFlowable, Image, KeepTogether, PageBreak,
                                Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

import action_plan

BLOCK_ORDER = ["summary", "symptoms", "cause", "reason", "precaution", "prevention"]


def advisory_ok(report):
    """True when a written advisory passed the audit.

    The downloads mirror the web page: the advisory replaces the raw passage blocks
    when it exists, and the blocks come back when it does not, so a report is never
    empty just because the API was down.
    """
    adv = report.get("advisory")
    return bool(adv and adv.get("ok"))

BAND_COLOUR = {"Healthy": "#2e7d32", "Mild": "#f9a825", "Moderate": "#ef6c00",
               "High": "#c62828", "Unavailable": "#546e7a"}


def _decode(data_uri):
    """`data:image/png;base64,AAAA...` -> raw bytes, or None."""
    if not data_uri or not isinstance(data_uri, str):
        return None
    m = re.match(r"data:image/[a-zA-Z+]+;base64,(.+)$", data_uri, re.S)
    if not m:
        return None
    try:
        return base64.b64decode(m.group(1))
    except Exception:                                  # noqa: BLE001
        return None


def _esc(text):
    """Escape for reportlab's mini-HTML paragraph markup."""
    return (str(text).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;"))


def _plan(report):
    return report.get("plan") or action_plan.build(report)


# ------------------------------------------------------------------ PDF
def to_pdf(report):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
        title=f"PlantCare report - {report.get('title', '')}",
        author="PlantCare-AI")

    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=13.5,
                          spaceAfter=5, alignment=TA_LEFT)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=19, leading=23,
                        spaceAfter=2, textColor=colors.HexColor("#14301c"))
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=11.5, leading=15,
                        spaceBefore=11, spaceAfter=4,
                        textColor=colors.HexColor("#2e7d32"))
    small = ParagraphStyle("small", parent=body, fontSize=8, leading=11,
                           textColor=colors.HexColor("#6b7280"))
    cite = ParagraphStyle("cite", parent=small, spaceAfter=8)
    step = ParagraphStyle("step", parent=body, leftIndent=13, spaceAfter=2)

    S = []
    S.append(Paragraph("PLANTCARE-AI &nbsp;·&nbsp; LEAF DIAGNOSIS REPORT", small))
    S.append(HRFlowable(width="100%", thickness=0.8,
                        color=colors.HexColor("#2e7d32"), spaceAfter=8))

    S.append(Paragraph(_esc(report.get("title", "Diagnosis")), h1))
    S.append(Paragraph(
        f"{_esc(report.get('crop', ''))} &nbsp;·&nbsp; "
        f"{_esc(report.get('confidence_pct', ''))} confidence "
        f"({_esc(report.get('confidence_band', ''))}) &nbsp;·&nbsp; "
        f"evidence tier {_esc(report.get('tier', '-'))}", small))
    S.append(Paragraph(
        f"{_esc(report.get('filename', ''))} &nbsp;·&nbsp; "
        f"{_esc(report.get('generated_at', ''))}", small))
    S.append(Spacer(1, 9))

    # severity band + image side by side
    band = report.get("severity_level", "Unavailable")
    sev_cell = [
        Paragraph(f"<b>Severity: {_esc(band)}</b>", body),
        Paragraph(_esc(report.get("severity_display") or "not measured")
                  + " of leaf area", small),
        Spacer(1, 3),
        Paragraph(_esc(report.get("severity_advice", "")), body),
    ]
    img_bytes = _decode((report.get("images") or {}).get("annotated"))
    if img_bytes:
        try:
            pic = Image(io.BytesIO(img_bytes), width=62 * mm, height=62 * mm,
                        kind="proportional")
            head = Table([[sev_cell, pic]], colWidths=[104 * mm, 66 * mm])
        except Exception:                              # noqa: BLE001
            head = Table([[sev_cell]], colWidths=[170 * mm])
    else:
        head = Table([[sev_cell]], colWidths=[170 * mm])
    head.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, 0), 8),
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#f4f7f4")),
        ("LINEBEFORE", (0, 0), (0, 0), 2.5, colors.HexColor(
            BAND_COLOUR.get(band, "#546e7a"))),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    S.append(head)

    # ---- action plan ----
    plan = _plan(report)
    S.append(Paragraph("What to do now", h2))
    S.append(Paragraph(f"<b>{_esc(plan['urgency'])}</b> — "
                       f"{_esc(plan['urgency_note'])}", body))
    for group in plan["groups"]:
        rows = [Paragraph(f"<b>{_esc(group['when'])}</b>", body)]
        for i, st in enumerate(group["steps"], 1):
            lbl = f"<b>{_esc(st['label'])}</b> — " if st["label"] else ""
            rows.append(Paragraph(f"{i}. {lbl}{_esc(st['text'])}", step))
            if st["citation"]:
                rows.append(Paragraph(_esc(st["citation"]), cite))
        S.append(KeepTogether(rows))
    for c in plan["closing"]:
        S.append(Paragraph(f"<i>{_esc(c)}</i>", small))

    # ---- evidence blocks ----
    sections = report.get("sections", {})
    adv = report.get("advisory")
    if advisory_ok(report):
        S.append(Paragraph("Field advisory", h2))
        for a in adv["sections"]:
            S.append(Paragraph(_esc(a["heading"]), h2))
            if a.get("body"):
                S.append(Paragraph(_esc(a["body"]), body))
            for b in a.get("bullets", []):
                S.append(Paragraph("&bull; " + _esc(b), body))

    for key in (BLOCK_ORDER if not advisory_ok(report) else []):
        sec = sections.get(key)
        if not sec:
            continue
        S.append(Paragraph(_esc(sec.get("heading", key.title())), h2))
        if sec.get("note"):
            S.append(Paragraph(f"<i>{_esc(sec['note'])}</i>", small))
        for p in sec.get("passages", []):
            S.append(Paragraph(_esc(p["text"]), body))
            S.append(Paragraph(_esc(p.get("citation", "")), cite))

    if report.get("caveats"):
        S.append(Paragraph("Read this before you act", h2))
        for c in report["caveats"]:
            S.append(Paragraph("• " + _esc(c), body))

    if report.get("citations"):
        S.append(Paragraph("Sources", h2))
        S.append(Paragraph(_esc("   ".join(report["citations"])), small))

    S.append(Spacer(1, 8))
    S.append(HRFlowable(width="100%", thickness=0.5,
                        color=colors.HexColor("#d4d9d4"), spaceAfter=5))
    S.append(Paragraph(
        "Guidance is quoted from published IPM documents; every line carries the "
        "document and page it came from. This report supports a decision — it does "
        "not replace an agriculture officer. Confirm with your nearest KVK before "
        "applying any chemical.", small))

    doc.build(S)
    buf.seek(0)
    return buf


# ------------------------------------------------------------------ DOCX
def to_docx(report):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    eyebrow = doc.add_paragraph()
    run = eyebrow.add_run("PLANTCARE-AI  ·  LEAF DIAGNOSIS REPORT")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_heading(report.get("title", "Diagnosis"), level=0)

    meta = doc.add_paragraph()
    run = meta.add_run(
        f"{report.get('crop', '')}  ·  {report.get('confidence_pct', '')} confidence "
        f"({report.get('confidence_band', '')})  ·  evidence tier "
        f"{report.get('tier', '-')}\n"
        f"{report.get('filename', '')}  ·  {report.get('generated_at', '')}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    img_bytes = _decode((report.get("images") or {}).get("annotated"))
    if img_bytes:
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(3.1))
        except Exception:                              # noqa: BLE001
            pass

    doc.add_heading(f"Severity: {report.get('severity_level', 'Unavailable')}", level=1)
    doc.add_paragraph(
        f"{report.get('severity_display') or 'not measured'} of leaf area · "
        f"{report.get('lesion_count', 0)} lesion patches")
    doc.add_paragraph(report.get("severity_advice", ""))

    plan = _plan(report)
    doc.add_heading("What to do now", level=1)
    lead = doc.add_paragraph()
    lead.add_run(plan["urgency"]).bold = True
    lead.add_run(" — " + plan["urgency_note"])

    for group in plan["groups"]:
        doc.add_heading(group["when"], level=2)
        for st in group["steps"]:
            para = doc.add_paragraph(style="List Number")
            if st["label"]:
                para.add_run(st["label"] + " — ").bold = True
            para.add_run(st["text"])
            if st["citation"]:
                cite = doc.add_paragraph()
                cite.paragraph_format.left_indent = Inches(0.4)
                run = cite.add_run(st["citation"])
                run.font.size = Pt(7.5)
                run.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    for c in plan["closing"]:
        para = doc.add_paragraph()
        run = para.add_run(c)
        run.italic = True
        run.font.size = Pt(8.5)

    sections = report.get("sections", {})

    adv = report.get("advisory")
    if advisory_ok(report):
        doc.add_heading("Field advisory", level=1)
        for a in adv["sections"]:
            doc.add_heading(a["heading"], level=2)
            if a.get("body"):
                doc.add_paragraph(a["body"])
            for b in a.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    for key in (BLOCK_ORDER if not advisory_ok(report) else []):
        sec = sections.get(key)
        if not sec:
            continue
        doc.add_heading(sec.get("heading", key.title()), level=1)
        if sec.get("note"):
            note = doc.add_paragraph()
            run = note.add_run(sec["note"])
            run.italic = True
            run.font.size = Pt(8.5)
        for p in sec.get("passages", []):
            doc.add_paragraph(p["text"])
            cite = doc.add_paragraph()
            run = cite.add_run(p.get("citation", ""))
            run.font.size = Pt(7.5)
            run.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    if report.get("caveats"):
        doc.add_heading("Read this before you act", level=1)
        for c in report["caveats"]:
            doc.add_paragraph(c, style="List Bullet")

    if report.get("citations"):
        doc.add_heading("Sources", level=1)
        doc.add_paragraph("   ".join(report["citations"]))

    footer = doc.add_paragraph()
    run = footer.add_run(
        "Guidance is quoted from published IPM documents; every line carries the "
        "document and page it came from. This report supports a decision — it does "
        "not replace an agriculture officer. Confirm with your nearest KVK before "
        "applying any chemical.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def filename(report, ext):
    base = re.sub(r"[^a-z0-9]+", "_", str(report.get("label", "report")).lower())
    return f"plantcare_{base.strip('_')}.{ext}"
